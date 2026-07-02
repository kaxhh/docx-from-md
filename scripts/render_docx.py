#!/usr/bin/env python3
"""JSON → DOCX renderer for AI+工业设计文档模板.

Derived from template: skill-test/AI+template.docx
Template analysis: skill-test/docx-template-analysis/

Supported table families:
  - macro_definition_table (3 col, no merges)
  - struct_member_table (3 col, no merges)
  - global_variable_table (3 col, no merges)
  - interface_spec_table (4 col, vMerge + gridSpan)

Fidelity: High — clones template styles, uses fixed DXA widths,
preserves heading numbering, table cell style '表格内文字',
caption style 'Caption', header fill D9D9D9, border sz=4.

Usage:
    python3 render_docx.py \
        --template template.docx \
        --input content_data.json \
        --output generated.docx
"""

import argparse
import copy
import json
import os
import re
import shutil
import sys

from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# ---------------------------------------------------------------------------
# Table configuration — derived from table_catalog.json
# ---------------------------------------------------------------------------
TABLE_CONFIG = {
    "macro_definition_table": {
        "grid_columns": [2305, 1736, 4341],
        "header_labels": ["宏名称", "宏内容", "说明"],
        "row_keys": ["name", "value", "description"],
    },
    "struct_member_table": {
        "grid_columns": [2683, 1991, 3818],
        "header_labels": ["成员变量", "类型", "说明"],
        "row_keys": ["name", "type", "description"],
    },
    "global_variable_table": {
        "grid_columns": [5270, 1674, 1548],
        "header_labels": ["全局变量名称", "类型", "说明"],
        "row_keys": ["name", "type", "description"],
    },
    "interface_spec_table": {
        "grid_columns": [1378, 1984, 1276, 4068],
        "row_labels": {
            "prototype": "接口原型",
            "description": "接口描述",
            "parameters": "接口参数",
            "returns": "返回值",
            "constraints": "使用限制",
            "notes": "其它说明",
        },
        "param_sub_header": ["数据类型", "参数名称", "参数说明"],
        "return_sub_header": ["数据类型", "返回值说明"],
    },
}

# Common formatting constants (from template analysis)
HEADER_FILL = "D9D9D9"
CELL_PARA_STYLE = "表格内文字"
CAPTION_STYLE = "Caption"
CAPTION_CHAR_STYLE = "题注 Char"
BORDER_SZ = "4"
ROW_HEIGHT_VAL = "270"
CELL_MARGIN_LEFT = 108
CELL_MARGIN_RIGHT = 108


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------
def _make_element(tag, **attrs):
    """Create an OxmlElement with attributes."""
    elem = OxmlElement(tag)
    for k, v in attrs.items():
        elem.set(qn(f"w:{k}"), str(v))
    return elem


def _set_cell_shading(tc, fill):
    """Set cell background fill."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    # Remove existing shd
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _set_cell_width(tc, width_dxa):
    """Set cell width in DXA."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(width_dxa))
    tcW.set(qn("w:type"), "dxa")


def _set_cell_borders(tc, sz=BORDER_SZ):
    """Set all four cell borders."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    for side in ("top", "left", "bottom", "right"):
        elem = borders.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:color"), "auto")
        elem.set(qn("w:sz"), sz)
        elem.set(qn("w:space"), "0")


def _set_cell_valign(tc, val="center"):
    """Set cell vertical alignment."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    vAlign = tcPr.find(qn("w:vAlign"))
    if vAlign is None:
        vAlign = OxmlElement("w:vAlign")
        tcPr.append(vAlign)
    vAlign.set(qn("w:val"), val)


def _set_grid_span(tc, span):
    """Set w:gridSpan on a cell."""
    if span <= 1:
        return
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    gs = tcPr.find(qn("w:gridSpan"))
    if gs is None:
        gs = OxmlElement("w:gridSpan")
        tcPr.append(gs)
    gs.set(qn("w:val"), str(span))


def _set_vmerge(tc, mode):
    """Set vMerge on a cell: 'restart' or 'continue'."""
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        vm = OxmlElement("w:vMerge")
        tcPr.append(vm)
    if mode == "restart":
        vm.set(qn("w:val"), "restart")
    # 'continue' = no w:val attribute (Word convention)


def _remove_grid_span_cells(tr, current_idx, span):
    """Remove (span-1) physical cells after the current cell."""
    if span <= 1:
        return
    tcs = list(tr.iterchildren(qn("w:tc")))
    to_remove = tcs[current_idx + 1: current_idx + span]
    for tc in to_remove:
        tr.remove(tc)


def _set_row_height(tr, val=ROW_HEIGHT_VAL, rule="atLeast"):
    """Set row height."""
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        tr.insert(0, trPr)
    trHeight = trPr.find(qn("w:trHeight"))
    if trHeight is None:
        trHeight = OxmlElement("w:trHeight")
        trPr.append(trHeight)
    trHeight.set(qn("w:val"), val)
    trHeight.set(qn("w:hRule"), rule)


def _set_table_borders(tbl, outer_sz="6", inner_sz=BORDER_SZ):
    """Set table-level borders with outer/inner distinction."""
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for side in ("top", "left", "bottom", "right"):
        elem = borders.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:color"), "auto")
        elem.set(qn("w:sz"), outer_sz)
        elem.set(qn("w:space"), "0")

    for side in ("insideH", "insideV"):
        elem = borders.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:color"), "auto")
        elem.set(qn("w:sz"), inner_sz)
        elem.set(qn("w:space"), "0")


def _set_table_cell_margins(tbl):
    """Set table cell margins."""
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        return
    mar = tblPr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tblPr.append(mar)
    for side, val in [("top", 0), ("left", CELL_MARGIN_LEFT), ("bottom", 0), ("right", CELL_MARGIN_RIGHT)]:
        elem = mar.find(qn(f"w:{side}"))
        if elem is None:
            elem = OxmlElement(f"w:{side}")
            mar.append(elem)
        elem.set(qn("w:w"), str(val))
        elem.set(qn("w:type"), "dxa")


def _write_cell_text(cell, text, style_name=CELL_PARA_STYLE):
    """Write text into a cell, splitting on newlines into separate paragraphs."""
    # Clear existing content
    for p in cell.paragraphs:
        p_elem = p._element
        for child in list(p_elem):
            if child.tag != qn("w:pPr"):
                p_elem.remove(child)

    lines = text.split("\n") if text else [""]
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.style = style_name if style_name else p.style
        run = p.add_run(line)


def _make_seq_field_paragraph(doc, prefix, seq_label, caption_text,
                               chapter_num="0", styleref_level=1):
    """Build a caption paragraph with SEQ field (auto-numbering).

    Produces XML like:
      "图 " [STYLEREF <level> \\s] [-] [SEQ 图 \\* ARABIC \\s 1] " " caption_text

    Args:
        doc: Document object (needed for paragraph creation)
        prefix: Caption prefix, e.g. "图" or "表"
        seq_label: SEQ identifier, e.g. "图" or "表"
        caption_text: The caption body text (after the number)
        chapter_num: Chapter number placeholder (before field update)
        styleref_level: Heading level for STYLEREF (e.g. 1 for "heading 1")

    Returns:
        The created paragraph element.
    """
    p = doc.add_paragraph()
    p.style = CAPTION_STYLE
    p_elem = p._element

    # Try to resolve the 题注 Char character style from the template
    char_style_id = None
    try:
        char_style = doc.styles[CAPTION_CHAR_STYLE]
        char_style_id = char_style.style_id
    except KeyError:
        pass

    def _rpr():
        """Create standard run properties for SEQ field runs."""
        rPr = OxmlElement('w:rPr')
        if char_style_id:
            rs = OxmlElement('w:rStyle')
            rs.set(qn('w:val'), char_style_id)
            rPr.append(rs)
        b = OxmlElement('w:b')
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(b)
        rPr.append(u)
        return rPr

    def _run_with_text(text, east_asia=False):
        """Add a run with text."""
        r = OxmlElement('w:r')
        r.append(_rpr())
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = text
        r.append(t)
        p_elem.append(r)

    def _fldchar(ftype):
        """Add a fldChar run."""
        r = OxmlElement('w:r')
        r.append(_rpr())
        fc = OxmlElement('w:fldChar')
        fc.set(qn('w:fldCharType'), ftype)
        r.append(fc)
        p_elem.append(r)

    def _instr_text(text, east_asia=False):
        """Add an instrText run."""
        r = OxmlElement('w:r')
        rPr = _rpr()
        if east_asia:
            rf = OxmlElement('w:rFonts')
            rf.set(qn('w:hint'), 'eastAsia')
            rPr.insert(0, rf)
        r.append(rPr)
        it = OxmlElement('w:instrText')
        it.set(qn('xml:space'), 'preserve')
        it.text = text
        r.append(it)
        p_elem.append(r)

    # 1. Prefix text (e.g. "图 ")
    _run_with_text(f'{prefix} ')

    # 2. STYLEREF field: { STYLEREF <level> \s } -> chapter number
    _fldchar('begin')
    _instr_text(f' STYLEREF {styleref_level} \\s ')
    _fldchar('separate')
    _run_with_text(chapter_num)  # placeholder: shows chapter number when field not updated
    _fldchar('end')

    # 3. noBreakHyphen separator
    r = OxmlElement('w:r')
    r.append(_rpr())
    nbh = OxmlElement('w:noBreakHyphen')
    r.append(nbh)
    p_elem.append(r)

    # 4. SEQ field: { SEQ 图 \* ARABIC \s 1 } -> sequence number
    _fldchar('begin')
    _instr_text(f' SEQ ')
    _instr_text(seq_label, east_asia=True)
    _instr_text(f' \\* ARABIC \\s 1 ')
    _fldchar('separate')
    _run_with_text('0')  # placeholder
    _fldchar('end')

    # 5. Space + caption body text
    _run_with_text(' ', east_asia=True)

    # Caption body text (with east-Asia font hint)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:hint'), 'eastAsia')
    rf.set(qn('w:eastAsia'), '楷体')
    rf.set(qn('w:cs'), '楷体')
    rPr.append(rf)
    kern = OxmlElement('w:kern')
    kern.set(qn('w:val'), '2')
    rPr.append(kern)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '24')
    rPr.append(szCs)
    lang = OxmlElement('w:lang')
    lang.set(qn('w:bidi'), 'ar')
    rPr.append(lang)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = caption_text
    r.append(t)
    p_elem.append(r)

    return p


# ---------------------------------------------------------------------------
# Table renderers
# ---------------------------------------------------------------------------
def render_simple_table(doc, table_data, config):
    """Render a 3-column flat table (macro/struct/global)."""
    grid_cols = config["grid_columns"]
    header_labels = config["header_labels"]
    row_keys = config["row_keys"]
    total_width = sum(grid_cols)

    num_rows = 1 + len(table_data.get("rows", []))
    table = doc.add_table(rows=num_rows, cols=len(grid_cols))

    # Table properties
    tbl_elem = table._tbl
    _set_table_borders(tbl_elem)
    _set_table_cell_margins(tbl_elem)

    # Set grid column widths
    tblGrid = tbl_elem.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl_elem.insert(1, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in grid_cols:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    # Header row
    header_row = table.rows[0]
    _set_row_height(header_row._tr)
    for col_idx, (label, width) in enumerate(zip(header_labels, grid_cols)):
        cell = header_row.cells[col_idx]
        tc = cell._tc
        _set_cell_width(tc, width)
        _set_cell_shading(tc, HEADER_FILL)
        _set_cell_borders(tc)
        _set_cell_valign(tc)
        _write_cell_text(cell, label)

    # Body rows
    for row_idx, row_data in enumerate(table_data.get("rows", []), start=1):
        row = table.rows[row_idx]
        _set_row_height(row._tr)
        for col_idx, (key, width) in enumerate(zip(row_keys, grid_cols)):
            cell = row.cells[col_idx]
            tc = cell._tc
            _set_cell_width(tc, width)
            _set_cell_shading(tc, "auto")
            _set_cell_borders(tc)
            _set_cell_valign(tc)
            _write_cell_text(cell, row_data.get(key, ""))

    return table


def render_interface_spec_table(doc, table_data, config):
    """Render a merged interface specification table."""
    grid_cols = config["grid_columns"]
    row_labels = config["row_labels"]
    total_width = sum(grid_cols)
    rows_data = table_data.get("rows", {})

    # Build row plan: list of (cells, merge_info)
    # Each cell: (text, span, is_header)
    row_plan = []

    # 1. 接口原型
    proto_text = rows_data.get("prototype", "")
    row_plan.append([
        (row_labels["prototype"], 1, True),
        (proto_text, 3, False),
    ])

    # 2. 接口描述
    desc_text = rows_data.get("description", "")
    row_plan.append([
        (row_labels["description"], 1, True),
        (desc_text, 3, False),
    ])

    # 3. 接口参数
    params = rows_data.get("parameters", [])
    if params:
        row_plan.append([
            (row_labels["parameters"], 1, True, "restart"),
            (config["param_sub_header"][0], 1, True),
            (config["param_sub_header"][1], 1, True),
            (config["param_sub_header"][2], 1, True),
        ])
        for p in params:
            row_plan.append([
                ("", 1, True, "continue"),
                (p.get("type", ""), 1, False),
                (p.get("name", ""), 1, False),
                (p.get("description", ""), 1, False),
            ])
    else:
        row_plan.append([
            (row_labels["parameters"], 1, True),
            ("无", 3, False),
        ])

    # 4. 返回值
    returns = rows_data.get("returns", [])
    if returns:
        row_plan.append([
            (row_labels["returns"], 1, True, "restart"),
            (config["return_sub_header"][0], 1, True),
            (config["return_sub_header"][1], 2, True),
        ])
        for r in returns:
            row_plan.append([
                ("", 1, True, "continue"),
                (r.get("type", ""), 1, False),
                (r.get("description", ""), 2, False),
            ])
    else:
        row_plan.append([
            (row_labels["returns"], 1, True),
            ("无", 3, False),
        ])

    # 5. 使用限制
    row_plan.append([
        (row_labels["constraints"], 1, True),
        (rows_data.get("constraints", "无"), 3, False),
    ])

    # 6. 其它说明
    row_plan.append([
        (row_labels["notes"], 1, True),
        (rows_data.get("notes", "无"), 3, False),
    ])

    # Create table
    num_rows = len(row_plan)
    table = doc.add_table(rows=num_rows, cols=4)

    tbl_elem = table._tbl
    _set_table_borders(tbl_elem)
    _set_table_cell_margins(tbl_elem)

    # Set grid column widths
    tblGrid = tbl_elem.find(qn("w:tblGrid"))
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl_elem.insert(1, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in grid_cols:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        tblGrid.append(gc)

    # Render rows
    for r_idx, cells_spec in enumerate(row_plan):
        tr = table.rows[r_idx]._tr
        _set_row_height(tr)
        physical_col = 0
        grid_col = 0

        for cell_spec in cells_spec:
            if len(cell_spec) == 4:
                text, span, is_header, vmerge_mode = cell_spec
            else:
                text, span, is_header = cell_spec
                vmerge_mode = None

            cell = table.rows[r_idx].cells[physical_col]
            tc = cell._tc

            # Width = sum of grid columns covered
            span_width = sum(grid_cols[grid_col + s] for s in range(span))
            _set_cell_width(tc, span_width)
            _set_cell_borders(tc)
            _set_cell_valign(tc)

            # Shading
            if is_header:
                _set_cell_shading(tc, HEADER_FILL)
            else:
                _set_cell_shading(tc, "auto")

            # GridSpan
            _set_grid_span(tc, span)

            # vMerge
            if vmerge_mode:
                _set_vmerge(tc, vmerge_mode)

            # Text
            _write_cell_text(cell, text)

            # Remove covered physical cells
            _remove_grid_span_cells(tr, physical_col, span)

            physical_col += 1
            grid_col += span

    return table


# ---------------------------------------------------------------------------
# Document renderer
# ---------------------------------------------------------------------------
class DocxRenderer:
    """Render JSON content into a DOCX document based on the template."""

    HEADING_STYLES = {
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        4: "Heading 4",
        5: "Heading 5",
        6: "Heading 6",
    }

    def __init__(self, template_path, base_dir="."):
        self.template_path = template_path
        self.base_dir = base_dir
        self.doc = Document(template_path)
        self._prepare_template()

    def _prepare_template(self):
        """Remove template body content, keeping styles/sections/headers/footers."""
        # Remove all body elements except sectPr (section properties)
        body = self.doc.element.body
        elements_to_remove = []
        for child in list(body):
            if child.tag == qn("w:sectPr"):
                continue
            elements_to_remove.append(child)
        for elem in elements_to_remove:
            body.remove(elem)

    def render(self, data, output_path):
        """Render JSON data to output .docx file."""
        sections = data.get("sections", [])

        # Detect minimum heading level for STYLEREF field in captions.
        # e.g. if top-level is ##, min_level=2 → STYLEREF 2 \s
        self._styleref_level = self._find_min_level(sections)

        # Extract chapter number from the first top-level section title
        # e.g., "4.7 分区文件系统" → chapter_num = "4"
        self._chapter_num = "0"
        for section in sections:
            title = section.get("title", "")
            match = re.match(r'^(\d+)', title)
            if match:
                self._chapter_num = match.group(1)
                break

        # Render sections recursively
        for section in sections:
            self._render_section(section)

        self.doc.save(output_path)

    @staticmethod
    def _find_min_level(sections):
        """Recursively find the minimum heading level in sections."""
        if not sections:
            return 99  # sentinel: empty list shouldn't affect parent
        min_lv = 99
        for s in sections:
            lv = s.get("level", 1)
            if lv < min_lv:
                min_lv = lv
            child_min = DocxRenderer._find_min_level(s.get("children", []))
            if child_min < min_lv:
                min_lv = child_min
        return min_lv if min_lv <= 6 else 1

    def _render_section(self, section):
        """Render a section (heading + content + children)."""
        level = section.get("level", 1)
        title = section.get("title", "")

        # Add heading with explicit outline level for TOC \u flag
        style_name = self.HEADING_STYLES.get(level, "Heading 1")
        heading_para = self.doc.add_heading(title, level=level)

        # Set outline level on paragraph (required by TOC \u)
        pPr = heading_para._element.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            heading_para._element.insert(0, pPr)
        outlineLvl = pPr.find(qn('w:outlineLvl'))
        if outlineLvl is None:
            outlineLvl = OxmlElement('w:outlineLvl')
            pPr.append(outlineLvl)
        outlineLvl.set(qn('w:val'), str(level - 1))

        # Add code blocks first (before paragraphs, so they don't land between caption and table)
        for block in section.get("code_blocks", []):
            self._render_code_block(block)

        # Add paragraphs
        for para in section.get("paragraphs", []):
            self._render_paragraph(para)

        # Add tables
        for table_data in section.get("tables", []):
            self._render_table(table_data)

        # Add figures
        for fig in section.get("figures", []):
            self._render_figure(fig)

        # Render children
        for child in section.get("children", []):
            self._render_section(child)

    def _render_paragraph(self, para):
        """Render a paragraph block."""
        text = para.get("text", "")
        para_type = para.get("type", "text")

        if para_type == "caption":
            # Detect caption kind and build SEQ field
            if text.startswith("图"):
                prefix, seq_label = "图", "图"
                body = text[len("图"):].lstrip()
            elif text.startswith("表"):
                prefix, seq_label = "表", "表"
                body = text[len("表"):].lstrip()
            elif text.startswith("代码"):
                prefix, seq_label = "代码", "代码"
                body = text[len("代码"):].lstrip()
            else:
                prefix, seq_label, body = text[:1], text[:1], text[1:].lstrip()

            _make_seq_field_paragraph(self.doc, prefix, seq_label, body,
                                       self._chapter_num, self._styleref_level)
        elif para_type == "text":
            self.doc.add_paragraph(text)

    def _render_code_block(self, block):
        """Render a code block."""
        code = block.get("code", "")
        lang = block.get("language", "")

        # Use '代码' style if available, else Normal with monospace
        p = self.doc.add_paragraph()
        try:
            p.style = "代码"
        except KeyError:
            pass

        run = p.add_run(code)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def _render_table(self, table_data):
        """Render a table based on its type."""
        table_type = table_data.get("type")
        config = TABLE_CONFIG.get(table_type)

        if config is None:
            print(f"WARNING: Unknown table type: {table_type}", file=sys.stderr)
            return

        # Caption is already rendered as a paragraph block — do NOT render
        # tables[].caption again (single source of truth rule)

        if table_type == "interface_spec_table":
            render_interface_spec_table(self.doc, table_data, config)
        else:
            render_simple_table(self.doc, table_data, config)

    def _render_figure(self, fig):
        """Render a figure with embedded image (SVG auto-converted to PNG)."""
        import os
        import tempfile
        import cairosvg

        source = fig.get("source", "")
        caption = fig.get("caption", "")

        # Resolve image path
        img_path = os.path.join(self.base_dir, source) if self.base_dir else source

        embedded = False
        if os.path.isfile(img_path):
            ext = os.path.splitext(img_path)[1].lower()
            if ext == '.svg':
                # Convert SVG to PNG and embed
                try:
                    # Read SVG and replace missing fonts
                    with open(img_path, 'r', encoding='utf-8') as f:
                        svg_content = f.read()
                    # Replace SimSun (宋体) with Noto Sans CJK SC
                    svg_content = svg_content.replace('SimSun', 'Noto Sans CJK SC')

                    # Write patched SVG to temp file
                    with tempfile.NamedTemporaryFile(suffix='.svg', delete=False, mode='w', encoding='utf-8') as tmp_svg:
                        tmp_svg.write(svg_content)
                        tmp_svg_path = tmp_svg.name

                    # Convert to PNG with higher DPI for better quality
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_png:
                        cairosvg.svg2png(url=tmp_svg_path, write_to=tmp_png.name, dpi=150)
                        self.doc.add_picture(tmp_png.name)
                        embedded = True
                        os.unlink(tmp_png.name)

                    os.unlink(tmp_svg_path)
                except Exception as e:
                    print(f"⚠️  SVG 转换失败 {source}: {e}", file=sys.stderr)
            else:
                try:
                    self.doc.add_picture(img_path)
                    embedded = True
                except Exception as e:
                    print(f"⚠️  图片嵌入失败 {source}: {e}", file=sys.stderr)

        if not embedded:
            # Add placeholder
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[Figure: {source}]")
            run.font.size = Pt(10)

        # Add caption with SEQ field
        if caption:
            if caption.startswith("图"):
                body = caption[len("图"):].lstrip()
            else:
                body = caption
            _make_seq_field_paragraph(self.doc, "图", "图", body,
                                       self._chapter_num, self._styleref_level)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(description="JSON → DOCX renderer for AI+设计文档模板")
    parser.add_argument("--template", required=True, help="Template .docx file")
    parser.add_argument("--input", required=True, help="Input JSON file")
    parser.add_argument("--output", required=True, help="Output .docx file")
    parser.add_argument("--base-dir", default=".", help="Base directory for resolving relative image paths")
    args = parser.parse_args()

    with open(args.input, "r", encoding="utf-8") as f:
        data = json.load(f)

    renderer = DocxRenderer(args.template, args.base_dir)
    renderer.render(data, args.output)

    print(f"Generated: {args.output}")


if __name__ == "__main__":
    main()
